#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_docx.py — 阶段三：按模板画像从 Markdown 稿生成可提交 .docx（hbst-thesis skill）

依据 references/03-template-reverse.md 的模板画像生成 6 节结构：
  封面(无页码) → 声明+授权书(各一页,无页码) → 中文摘要(罗马i) → ABSTRACT(罗马续)
  → 目录(罗马续, TOC域) → 正文~附录(阿拉伯1起, 页眉"河北科技师范学院学士学位论文")

用法:
  python build_docx.py manuscript.md --out thesis.docx [--img-max-width-cm 14]

manuscript.md 约定（front matter + 顺序章节）:
```markdown
---
school: 河北科技师范学院
doc_type: 学士学位论文
title: 基于……的设计与实现
title_en: Design and Implementation of ...
name: 孟凡宇
student_id: 0963220118
college: 数学与信息科技学院
major: 网络工程
advisor: 蔡黔鹰
date: 二〇二六年五月二十日
cover_logo: 封面校徽图片路径(可选)
---
# 摘　　要
摘要正文……（首行缩进由脚本处理；[1] 会被转为上标引用）
# ABSTRACT
英文摘要……
# 目　　录
<!--TOC-->        ← 该标记处插入自动目录域(Word 中右键更新)
# 第1章　绪论
## 1.1　研究背景
正文段落。参考文献 [1] 与 [2] 是上标。
表4.1　用户表(user)      ← 表格上一行以"表"开头 = 表题注(表上方)
| 编号 | 字段名 | 类型 | 注释 |
| --- | --- | --- | --- |
| 1 | id | bigint | 用户ID |
![图4.1　系统架构图](figs/arch.png)   ← 图片: alt为题注(图下方)
# 参考文献
[1] 作者. 标题[J]. 刊名, 年, 卷(期): 页.
# 致　　谢
……
# 附　　录
……
```

依赖: python-docx (pip install python-docx)；PIL 可选(用于按宽高自动缩放图片)。
生成后请用 check_docx_structure.py 检查，并导出 PDF 核对版面(writing-standards §10)。
"""
import argparse
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# ---------- 模板常量（河北科技师范学院画像，见 03-template-reverse.md） ----------
MARGIN = dict(top=Cm(2.5), bottom=Cm(2.0), left=Cm(2.0), right=Cm(2.0))
GUTTER_TWIPS = 284
BODY_FONT_CN = "宋体"
BODY_FONT_EN = "Times New Roman"
HEADING_FONT_CN = "黑体"
BODY_SIZE_PT = 12          # 小四
LINE_EXACT_PT = 20         # 固定 20 磅
FIRST_LINE_CHARS = 2       # 首行缩进 2 字符
HEADER_TEXT = "河北科技师范学院学士学位论文"

DECLARATION_TITLE = "学位论文原创性声明"
AUTHORIZATION_TITLE = "学位论文版权使用授权书"


# ---------- OOXML 辅助 ----------
def set_run(run, cn=BODY_FONT_CN, en=BODY_FONT_EN, size=BODY_SIZE_PT, bold=False, superscript=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    # 论文正文与标题一律纯黑（模板要求；避免继承 Word 内置 Heading 的蓝色主题色）
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    if superscript:
        run.font.superscript = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)
    rFonts.set(qn("w:eastAsia"), cn)


def set_para(p, align=None, indent_chars=0, line_exact=None, before=None, after=None):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    if indent_chars:
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:firstLineChars"), str(indent_chars * 100))
        ind.set(qn("w:firstLine"), str(int(indent_chars * BODY_SIZE_PT * 20)))
    if line_exact:
        pf.line_spacing_rule = 3  # WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_exact)
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)


def add_text_with_citations(p, text, cn=BODY_FONT_CN, en=BODY_FONT_EN, size=BODY_SIZE_PT, bold=False):
    """把 [n] 引用拆成上标 run，其余为正文 run。"""
    parts = re.split(r"(\[\d+(?:[-–]\d+)?\])", text)
    for part in parts:
        if not part:
            continue
        if re.fullmatch(r"\[\d+(?:[-–]\d+)?\]", part):
            r = p.add_run(part)
            set_run(r, cn=cn, en=en, size=size, bold=bold, superscript=True)
        else:
            r = p.add_run(part)
            set_run(r, cn=cn, en=en, size=size, bold=bold)


def set_section(section, start_type=None):
    if start_type is not None:
        section.start_type = start_type
    section.top_margin = MARGIN["top"]
    section.bottom_margin = MARGIN["bottom"]
    section.left_margin = MARGIN["left"]
    section.right_margin = MARGIN["right"]
    from docx.shared import Twips
    section.gutter = Twips(GUTTER_TWIPS)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def norm(t):
    """去全角/半角空格，用于标题关键词判断（如 目　　录 -> 目录）。"""
    return t.replace("\u3000", "").replace(" ", "")


def new_section_with_footer(doc, fmt, start=None, header=None):
    """新增节并设置页脚 PAGE 域（可带页眉）。返回新节。"""
    s = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section(s)
    set_pgnum(s, fmt=fmt, start=start)
    s.footer.is_linked_to_previous = False
    fp = s.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)
    if header:
        s.header.is_linked_to_previous = False
        hp = s.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(hp.add_run(header), size=9)
    return s





def set_pgnum(section, fmt=None, start=None):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def _fld(run, fldchar_type=None, val=None):
    """正确生成 w:fldChar / w:instrText 元素。
    fldChar 必须用 w:fldCharType 属性，Word 才能识别为域。"""
    el = OxmlElement("w:fldChar")
    if fldchar_type:
        el.set(qn("w:fldCharType"), fldchar_type)
    if val is not None:
        el.set(qn("w:val"), val)
    run._element.append(el)


def add_page_field(paragraph):
    r = paragraph.add_run()
    _fld(r, "begin")
    r = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    r._element.append(it)
    r = paragraph.add_run()
    _fld(r, "separate")
    r = paragraph.add_run("1")
    set_run(r, size=9)
    r = paragraph.add_run()
    _fld(r, "end")


def add_toc_field(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    _fld(r, "begin")
    r = p.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = ' TOC \\o "1-3" \\h \\z \\u '
    r._element.append(it)
    r = p.add_run()
    _fld(r, "separate")
    r = p.add_run("（目录为自动域：请在 Word 中全选后按 F9 / 右键“更新域”生成）")
    set_run(r, size=10.5)
    from docx.shared import RGBColor
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    r = p.add_run()
    _fld(r, "end")


def enable_update_fields(doc):
    settings = doc.settings.element
    uf = settings.find(qn("w:updateFields"))
    if uf is None:
        uf = OxmlElement("w:updateFields")
        settings.append(uf)
    uf.set(qn("w:val"), "true")


# ---------- 封面 / 声明 ----------
def build_cover(doc, meta):
    for _ in range(2):
        doc.add_paragraph()
    logo = meta.get("cover_logo")
    if logo and os.path.exists(logo):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo, width=Cm(8))
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(meta.get("doc_type", "学士学位论文")), cn="黑体", size=26, bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(meta.get("title", "")), cn="黑体", size=22, bold=True)
    if meta.get("title_en"):
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(meta["title_en"]), cn="宋体", en="Times New Roman", size=16)
    for _ in range(3):
        doc.add_paragraph()
    for label, key in (("姓    名", "name"), ("学    号", "student_id"),
                       ("院    系", "college"), ("专    业", "major"),
                       ("指导教师", "advisor")):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(f"{label}    {meta.get(key, '')}"), size=16)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(meta.get("date", "")), size=16)


def build_declarations(doc, meta):
    title = meta.get("title", "")
    # 页 1：原创性声明
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(DECLARATION_TITLE), cn="黑体", size=16, bold=True)
    body1 = (f"本人所提交的学位论文《{title}》，是在导师的指导下，独立进行研究工作所取得的原创性成果。"
             "除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已经发表或撰写过的研究成果。"
             "对本文的研究做出重要贡献的个人和集体，均已在文中标明。")
    body2 = "本声明的法律后果由本人承担。"
    for txt in (body1, body2):
        p = doc.add_paragraph(); set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_chars=2, line_exact=LINE_EXACT_PT)
        set_run(p.add_run(txt))
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_run(p.add_run("论文作者（签名）：                    指导教师确认（签名）："), size=14)
    p = doc.add_paragraph()
    set_run(p.add_run("年    月    日                          年    月    日"), size=14)
    doc.add_page_break()
    # 页 2：版权授权书
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(AUTHORIZATION_TITLE), cn="黑体", size=16, bold=True)
    auth1 = ("本学位论文作者完全了解河北科技师范学院有权保留并向国家有关部门或机构送交学位论文的复印件和磁盘，"
             "允许论文被查阅和借阅。本人授权河北科技师范学院可以将学位论文的全部或部分内容编入有关数据库进行检索，"
             "可以采用影印、缩印或其它复制手段保存、汇编学位论文。")
    auth2 = "保密的学位论文在_______年解密后适用本授权书。"
    for txt in (auth1, auth2):
        p = doc.add_paragraph(); set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_chars=2, line_exact=LINE_EXACT_PT)
        set_run(p.add_run(txt))
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_run(p.add_run("论文作者（签名）：                    指导教师（签名）："), size=14)
    p = doc.add_paragraph()
    set_run(p.add_run("年    月    日                          年    月    日"), size=14)


# ---------- Markdown 解析与正文生成 ----------
HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")
IMG_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def parse_front_matter(lines):
    meta = {}
    if not lines or not lines[0].strip().startswith("---"):
        return meta, lines
    end = None
    for i in range(1, len(lines)):
        if lines[i].strip() == "---":
            end = i
            break
    if end is None:
        return meta, lines
    for line in lines[1:end]:
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if ":" in line:
            k, v = line.split(":", 1)
            meta[k.strip()] = v.strip().strip('"').strip("'")
    return meta, lines[end + 1:]


def iter_blocks(lines):
    """产出 block: ('h', level, text) / ('p', text) / ('table', rows) / ('img', alt, path) /
    ('toc',) / ('ref', text) / ('pagebreak',)。参考文献区由调用方维护状态。"""
    i = 0
    n = len(lines)
    in_ref = False
    while i < n:
        raw = lines[i].rstrip("\n")
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "<!--TOC-->" or line == "<!-- toc -->":
            yield ("toc",)
            i += 1
            continue
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            in_ref = text.startswith("参考文献") or "参考文献" in text
            yield ("h", level, text)
            i += 1
            continue
        # 表格块（含表头分隔行）
        if line.startswith("|") and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1]):
            rows = []
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            yield ("table", header, rows)
            continue
        # 图片行
        mimg = IMG_RE.match(line)
        if mimg:
            yield ("img", mimg.group(1), mimg.group(2))
            i += 1
            continue
        # 参考文献条目（[n] 开头）
        if in_ref and re.match(r"^\[\d+\]\s*\S", line):
            yield ("ref", line)
            i += 1
            continue
        if in_ref:
            # 参考文献区内的续行并入上一条
            yield ("ref_cont", line)
            i += 1
            continue
        yield ("p", line)
        i += 1


def add_table(doc, header, rows):
    n_cols = max(len(header), *(len(r) for r in rows)) if rows else len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    # 填充
    all_rows = [header] + rows
    for ri, rowdata in enumerate(all_rows):
        for ci in range(n_cols):
            cell = table.cell(ri, ci)
            txt = rowdata[ci] if ci < len(rowdata) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(txt), size=10.5, bold=(ri == 0))
    # 三线表边框：清默认，设表级 top/bottom；表头行单元格底边框 = 栏目线
    tbl = table._tbl
    tblPr = tbl.tblPr
    # 移除样式边框
    for el in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(el)
    borders = OxmlElement("w:tblBorders")
    for edge, sz, val in (("top", "12", "single"), ("bottom", "12", "single"),
                          ("left", "", "nil"), ("right", "", "nil"),
                          ("insideH", "", "nil"), ("insideV", "", "nil")):
        e = OxmlElement(f"w:{edge}")
        if sz:
            e.set(qn("w:val"), val); e.set(qn("w:sz"), sz)
            e.set(qn("w:space"), "0"); e.set(qn("w:color"), "000000")
        else:
            e.set(qn("w:val"), val)
        borders.append(e)
    tblPr.append(borders)
    # 表头行单元格下边框（栏目线 0.75pt=sz6）
    hdr_tr = table.rows[0]._tr
    for tc in hdr_tr.findall(qn("w:tc")):
        tcPr = tc.get_or_add_tcPr()
        tcB = OxmlElement("w:tcBorders")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000")
        tcB.append(b)
        tcPr.append(tcB)


def add_heading(doc, level, text):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    style = style_map.get(level, "Heading 3")
    p = doc.add_paragraph(style=style)
    size = {1: 15, 2: 14, 3: 13}[level]
    bold = level in (1, 3)
    add_text_with_citations(p, text, cn=HEADING_FONT_CN, size=size, bold=bold)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 每个一级标题（章/结论/参考文献/致谢/附录）另起一页；
        # 分节换页后标题已在页首时该属性不产生额外空白页
        p.paragraph_format.page_break_before = True
    # 行距与段距（按模板 Heading 样式，显式兜底）
    set_para(p, line_exact=None)
    return p


def add_image(doc, alt, path, img_max_cm):
    if not os.path.exists(path):
        p = doc.add_paragraph()
        set_run(p.add_run(f"[缺图: {path}]"), size=10.5)
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 自动缩放：宽<=max, 高<=18cm，按比例
    width = Cm(img_max_cm)
    try:
        from PIL import Image
        im = Image.open(path)
        w, h = im.size
        max_h_cm = 18.0
        scale = min(img_max_cm / (w or 1) * 1.0, max_h_cm / (h or 1) * 1.0)
        width = Cm(min(img_max_cm, max(1.0, (w * scale))))
        if (w * scale) > img_max_cm:
            width = Cm(img_max_cm)
        if (h * scale) > max_h_cm:
            width = Cm(max_h_cm * w / h)
    except Exception:
        pass
    p.add_run().add_picture(path, width=width)
    # 图注（alt 或下一段以"图"开头）：黑体 11pt（模板"图标题q"），图下方
    caption = alt if alt.strip() else None
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(cp, before=0, after=6)
        set_run(cp.add_run(caption), cn=HEADING_FONT_CN, size=11)


def build_docx(manuscript_path, out_path, img_max_cm):
    with open(manuscript_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()
    meta, rest = parse_front_matter(lines)

    doc = Document()
    # 样式兜底：全文字体黑色（标题/正文均不得出现主题蓝），西文 Times New Roman
    from docx.shared import RGBColor
    for sname in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle"):
        try:
            st = doc.styles[sname]
        except KeyError:
            continue
        st.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT_EN
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)
    # 标题字体显式设黑体/字号（Heading 样式默认西文 Calibri Light、带蓝色）
    for sname, (cn, pt) in (("Heading 1", (HEADING_FONT_CN, 15)), ("Heading 2", (HEADING_FONT_CN, 14)),
                            ("Heading 3", (HEADING_FONT_CN, 13))):
        try:
            st = doc.styles[sname]
        except KeyError:
            continue
        st.font.name = BODY_FONT_EN
        st.font.size = Pt(pt)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), cn)
        st.font.bold = (sname != "Heading 2")  # 一级/三级加粗，二级不加粗（模板）
    # 正文段落级行距与缩进由各 add_* 显式控制，此处只兜底字体

    # 节 1：封面
    sec0 = doc.sections[0]
    set_section(sec0)
    build_cover(doc, meta)

    # 节 2：声明 + 授权书（无页码）
    sec1 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section(sec1)
    build_declarations(doc, meta)

    # 后续状态机
    body_started = False

    blocks = list(iter_blocks(rest))
    idx = 0
    pending_table_caption = None  # 表题注先出现，紧随其后是表格
    while idx < len(blocks):
        blk = blocks[idx]
        kind = blk[0]
        if kind == "h":
            level, text = blk[1], blk[2]
            nt = norm(text)
            if level == 1:
                if nt.startswith("摘") and not body_started:
                    new_section_with_footer(doc, "upperRoman", start=1)
                elif nt.upper().startswith("ABSTRACT"):
                    new_section_with_footer(doc, "upperRoman")
                elif "目录" in nt:
                    new_section_with_footer(doc, "upperRoman")
                elif nt.startswith("第") and "章" in nt and not body_started:
                    new_section_with_footer(doc, "decimal", start=1, header=HEADER_TEXT)
                    body_started = True
                # 正文节内的后续一级标题(结论/参考文献/致谢/附录)留在同节
            add_heading(doc, level, text)
        elif kind == "toc":
            # 前文若已有“目录”一级标题则已切节；否则在此补节+标题
            has_toc_heading = any("目录" in norm(b[2]) for b in blocks[:idx] if b[0] == "h" and b[1] == 1)
            if not has_toc_heading:
                new_section_with_footer(doc, "upperRoman")
                add_heading(doc, 1, "目　　录")
            add_toc_field(doc)
        elif kind == "p":
            text = blk[1]
            if text.startswith("表") and re.match(r"^表\s*\d+\.\d+", text):
                pending_table_caption = text
                idx += 1
                continue
            p = doc.add_paragraph()
            set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_chars=FIRST_LINE_CHARS, line_exact=LINE_EXACT_PT)
            add_text_with_citations(p, text)
        elif kind == "table":
            header, rows = blk[1], blk[2]
            if pending_table_caption:
                cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_para(cp, line_exact=None, after=3)
                cp.paragraph_format.keep_with_next = True
                set_run(cp.add_run(pending_table_caption), cn=HEADING_FONT_CN, size=11, bold=False)
                pending_table_caption = None
            add_table(doc, header, rows)
        elif kind == "img":
            add_image(doc, blk[1], blk[2], img_max_cm)
        elif kind == "ref":
            p = doc.add_paragraph()
            set_para(p, line_exact=LINE_EXACT_PT)
            p.paragraph_format.left_indent = Pt(21)
            p.paragraph_format.first_line_indent = Pt(-21)
            set_run(p.add_run(blk[1]), size=10.5)
        elif kind == "ref_cont":
            # 参考文献换行续文：追加到上一段
            if doc.paragraphs:
                doc.paragraphs[-1].add_run(blk[1])
        idx += 1

    enable_update_fields(doc)
    doc.save(out_path)
    print(f"[build_docx] wrote {out_path}")
    print(f"[build_docx] 提示: 用 Word/WPS 打开后全选按 F9 更新目录与页码域；再导出 PDF 做版面核对。")
    print(f"[build_docx] 封面与声明页为脚本生成的标准版；若学校要求官方封面/声明版式，请在 Word 中替换对应页面。")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("manuscript")
    ap.add_argument("--out", default="thesis.docx")
    ap.add_argument("--img-max-width-cm", type=float, default=14.0)
    args = ap.parse_args()
    build_docx(args.manuscript, args.out, args.img_max_width_cm)


if __name__ == "__main__":
    main()
