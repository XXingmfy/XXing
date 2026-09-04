#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
check_text_quality.py — 阶段三：文本质量闸门（hbst-thesis skill）

输入论文正文纯文本（.txt/.md，或直接给 .docx 时尝试用 python-docx 抽文本），
逐项检查并输出 JSON 报告 + 退出码（存在 FAIL 项则 exit 1）。

检查项（对应 writing-standards 与 humanize 自检的可脚本化部分）:
  T1 未完成措辞:  待补充 / 待补 / 初稿 / TBD / TODO / 后续补充 / 占位 / needs_user_screenshot
  T2 过程/证据措辞: 根据源码 / 从代码证据 / 当前材料 / 根据 README / 本文不编造（正文泄漏）
  T3 重复标点:     。。 ，， 、、 ；； ：： ？？ ！！
  T4 疑似错别/赘字: 页面页面 订单订单 系统系统 之类相邻重复词（二字以上重复）
  T5 超长句:       连续无标点超过阈值（默认 80 字，humanize B1 判 40 字，这里宽松给长文检测用，可 --max-sentence）
  T6 模板链高频:   综上所述 / 总而言之 / 众所周知 / 值得注意的是 / 不难发现 / 随着…的发展 计数
  T7 匿名引用:     有研究指出 / 相关研究表明 / 研究表明 / 据报道 / 有学者指出（应具名）
  T8 中文字数:     CJK 字符数（默认阈值 8000，--min-cjk；可验收级 10000）
  T9 引用闭合:     正文 [n] 最大编号 与 参考文献条目数（粗略，供人工复核）

用法:
  python check_text_quality.py <paper.txt|paper.md|paper.docx> [--min-cjk 8000] [--json out.json]
"""
import argparse
import json
import os
import re
import sys

FORBIDDEN_PLACEHOLDER = [
    "待补充", "待补", "初稿", "TBD", "TODO", "待完善", "后续补充",
    "占位", "needs_user_screenshot", "pending_user_screenshot", "（初稿）", "(初稿)",
]
FORBIDDEN_PROCESS = [
    "根据源码", "根据项目源码", "从代码证据", "根据代码", "当前材料",
    "根据 README", "根据 PRD", "本文不编造", "根据现有代码",
]
CLICHES = ["综上所述", "总而言之", "众所周知", "值得注意的是", "不难发现",
           "随着我国", "随着科技的", "在当今社会", "随着互联网的飞速发展", "近年来，随着"]
ANON_CITE = ["有研究指出", "相关研究表明", "研究表明", "据报道", "有学者指出",
             "国外研究表明", "国内研究表明", "专家指出", "调查显示"]


def extract_text(path):
    if path.lower().endswith(".docx"):
        try:
            from docx import Document
        except ImportError:
            print(json.dumps({"error": "python-docx needed for .docx input"}, ensure_ascii=False))
            sys.exit(2)
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for tb in doc.tables:
            for row in tb.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("--min-cjk", type=int, default=8000, help="CJK 字符数下限")
    ap.add_argument("--max-sentence", type=int, default=80, help="无标点连续长度阈值")
    ap.add_argument("--json", default=None, help="报告输出路径")
    args = ap.parse_args()

    text = extract_text(args.input)
    text_clean = re.sub(r"```.*?```", "", text, flags=re.S)  # 去掉代码块（若有）
    # 去掉含样式标签/序号前缀的抽取行（如 "[44] (论文正文q) xxx"），只保留正文
    text_clean = re.sub(r"^\s*\[\d+\]\s*\([^)]*\)\s*", "", text_clean, flags=re.M)
    text_clean = re.sub(r"^\s*\[\d+\]\s*", "", text_clean, flags=re.M)

    issues = []

    def add(code, level, msg):
        issues.append({"id": code, "level": level, "message": msg})

    # T1 未完成措辞
    for kw in FORBIDDEN_PLACEHOLDER:
        if kw in text_clean:
            add("T1", "FAIL", f"发现未完成措辞: {kw!r}（次数 {text_clean.count(kw)}）")
    # T2 过程/证据措辞
    for kw in FORBIDDEN_PROCESS:
        if kw in text_clean:
            add("T2", "FAIL", f"正文出现过程/证据措辞: {kw!r}（次数 {text_clean.count(kw)}）")
    # T3 重复标点
    for pat in ["。。", "，，", "、、", "；；", "：：", "？？", "！！"]:
        n = text_clean.count(pat)
        if n:
            add("T3", "FAIL", f"重复标点 {pat!r} ×{n}")
    # T4 相邻重复词
    dup = set(re.findall(r"([\u4e00-\u9fff]{2,4})\1", text_clean))
    # 排除合法词如"慢慢/刚刚/人人"由 {2,4} 不会命中单字；命中如"页面页面"
    for d in dup:
        add("T4", "WARN", f"疑似重复词: {d}{d}")
    # T5 超长句（只统计含 CJK 的片段，避免英文长串误报）
    for m in re.finditer(r"[^，。；：！？\n]{,%d}[^，。；：！？\n]" % (args.max_sentence + 1), text_clean):
        seg = m.group(0)
        if len(seg) > args.max_sentence and re.search(r"[\u4e00-\u9fff]", seg):
            add("T5", "WARN", f"超长无停顿片段 {len(seg)} 字: {seg[:40]}…")
            if len(seg) > args.max_sentence + 40:
                break  # 只报前若干
    # T6 模板链高频
    for kw in CLICHES:
        n = text_clean.count(kw)
        if n:
            add("T6", "WARN" if n <= 2 else "FAIL", f"模板化措辞 {kw!r} ×{n}")
    # T7 匿名引用
    for kw in ANON_CITE:
        if kw in text_clean:
            add("T7", "WARN", f"匿名引用 {kw!r}（humanize A3 要求具名，缺来源标【待作者补充】）")
    # T8 字数
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text_clean))
    if cjk < args.min_cjk:
        add("T8", "FAIL", f"CJK 字符数 {cjk} < 下限 {args.min_cjk}")
    # T9 引用闭合（粗略）
    cites = sorted({int(x) for x in re.findall(r"\[(\d+)\]", text_clean)})
    ref_lines = [l for l in text_clean.splitlines() if re.match(r"^\s*\[\d+\]\s*\S", l)]
    max_cite = max(cites) if cites else 0
    ref_count = len(ref_lines)
    if max_cite and max_cite > ref_count:
        add("T9", "FAIL", f"正文引用最大编号 [{max_cite}] > 参考文献条目数 {ref_count}（引用未闭合）")
    elif max_cite and ref_count and max_cite < ref_count:
        add("T9", "WARN", f"正文最大引用 [{max_cite}] < 参考文献条目 {ref_count}（可能有未被引用的文献）")

    # T10 三线表"注释"列超长（writing-standards §6.5：≤4 汉字概括，禁止枚举/长说明）
    table_rows = re.findall(r"^\s*\|[^\n]+$", text_clean, flags=re.M)
    long_note = []
    for row in table_rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        if len(cells) >= 7:  # 数据库字段表：编号|字段名|类型|长度|是否非空|是否主键|注释
            note = cells[6]
            if not re.fullmatch(r"[\s|:\-]*", note):
                # 汉字数 + 连续西文/数字段(每段算1字)
                han = len(re.findall(r"[\u4e00-\u9fff]", note))
                latingroups = len(re.findall(r"[A-Za-z0-9_./-]+", note))
                if han + latingroups > 4 and note not in ("编号", "注释", "字段名", "类型", "长度", "是否非空", "是否主键"):
                    long_note.append(note)
    if long_note:
        # 去重、限量展示
        uniq = []
        for x in long_note:
            if x not in uniq:
                uniq.append(x)
        add("T10", "FAIL", f"表'注释'列存在超4字条目 ×{len(long_note)}: {uniq[:8]}（应≤4汉字概括，见 writing-standards §6.5）")

    failed = [i for i in issues if i["level"] == "FAIL"]
    report = {
        "cjk_chars": cjk,
        "issues": issues,
        "failed_count": len(failed),
        "passed": len(failed) == 0,
        "note": "T7 匿名引用、T9 闭合、T10 注释长度为辅助信号，最终以人工复核为准；T5 阈值按 humanize B1 建议 40 字可加严。",
    }
    out = json.dumps(report, ensure_ascii=False, indent=2)
    if args.json:
        with open(args.json, "w", encoding="utf-8") as f:
            f.write(out)
    print(out)
    sys.exit(1 if failed else 0)


if __name__ == "__main__":
    main()
